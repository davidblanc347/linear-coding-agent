"""Extract structured content from Microsoft Word documents (.docx).

This module provides functionality to extract text, headings, images, and metadata
from Word documents using python-docx. The extracted content is structured to be
compatible with the existing RAG pipeline (LLM processing and Weaviate ingestion).

Example:
    Extract content from a Word document:

        from pathlib import Path
        from utils.word_processor import extract_word_content

        result = extract_word_content(Path("document.docx"))
        print(f"Extracted {len(result['paragraphs'])} paragraphs")
        print(f"Found {len(result['headings'])} headings")

    Extract only metadata:

        metadata = extract_word_metadata(Path("document.docx"))
        print(f"Title: {metadata['title']}")
        print(f"Author: {metadata['author']}")

Note:
    Requires python-docx library: pip install python-docx>=0.8.11
"""

from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from datetime import datetime
import io
import re

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    raise ImportError(
        "python-docx library is required for Word processing. "
        "Install with: pip install python-docx>=0.8.11"
    )

from utils.types import TOCEntry


def extract_word_metadata(docx_path: Path) -> Dict[str, Any]:
    """Extract metadata from Word document core properties.

    Reads the document's core properties (title, author, created date, etc.)
    and attempts to extract additional metadata from the first few paragraphs
    if core properties are missing.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        Dictionary containing metadata fields:
        - title (str): Document title
        - author (str): Document author
        - created (datetime): Creation date
        - modified (datetime): Last modified date
        - language (str): Document language (if available)
        - edition (str): Edition info (if found in content)

    Example:
        >>> metadata = extract_word_metadata(Path("doc.docx"))
        >>> print(metadata["title"])
        'On the Origin of Species'
    """
    doc = Document(docx_path)
    core_props = doc.core_properties

    metadata = {
        "title": core_props.title or "",
        "author": core_props.author or "",
        "created": core_props.created,
        "modified": core_props.modified,
        "language": "",
        "edition": "",
    }

    # If metadata missing, try to extract from first paragraphs
    # Common pattern: "TITRE: ...", "AUTEUR: ...", "EDITION: ..."
    if not metadata["title"] or not metadata["author"]:
        for para in doc.paragraphs[:10]:  # Check first 10 paragraphs
            text = para.text.strip()

            # Match patterns like "TITRE : On the Origin..."
            if text.upper().startswith("TITRE") and ":" in text:
                metadata["title"] = text.split(":", 1)[1].strip()

            # Match patterns like "AUTEUR Charles DARWIN"
            elif text.upper().startswith("AUTEUR") and ":" in text:
                metadata["author"] = text.split(":", 1)[1].strip()
            elif text.upper().startswith("AUTEUR "):
                metadata["author"] = text[7:].strip()  # Remove "AUTEUR "

            # Match patterns like "EDITION : Sixth London Edition..."
            elif text.upper().startswith("EDITION") and ":" in text:
                metadata["edition"] = text.split(":", 1)[1].strip()

    return metadata


def _get_heading_level(style_name: str) -> Optional[int]:
    """Extract heading level from Word style name.

    Args:
        style_name: Word paragraph style name (e.g., "Heading 1", "Heading 2").

    Returns:
        Heading level (1-9) if it's a heading style, None otherwise.

    Example:
        >>> _get_heading_level("Heading 1")
        1
        >>> _get_heading_level("Heading 3")
        3
        >>> _get_heading_level("Normal")
        None
    """
    # Match patterns: "Heading 1", "Heading 2", etc.
    match = re.match(r"Heading (\d)", style_name)
    if match:
        level = int(match.group(1))
        return level if 1 <= level <= 9 else None
    return None


def extract_word_images(
    doc: Document,
    output_dir: Path,
    doc_name: str,
) -> List[Path]:
    """Extract inline images from Word document.

    Saves all inline images (shapes, pictures) to the output directory
    with sequential numbering.

    Args:
        doc: python-docx Document object.
        output_dir: Directory to save extracted images.
        doc_name: Document name for image filename prefix.

    Returns:
        List of paths to extracted image files.

    Example:
        >>> doc = Document("doc.docx")
        >>> images = extract_word_images(doc, Path("output"), "darwin")
        >>> print(f"Extracted {len(images)} images")
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    image_paths: List[Path] = []

    image_counter = 0

    # Extract images from document relationships
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob

                # Determine file extension from content type
                content_type = rel.target_part.content_type
                ext = "png"  # default
                if "jpeg" in content_type or "jpg" in content_type:
                    ext = "jpg"
                elif "png" in content_type:
                    ext = "png"
                elif "gif" in content_type:
                    ext = "gif"

                # Save image
                image_filename = f"{doc_name}_image_{image_counter}.{ext}"
                image_path = output_dir / image_filename

                with open(image_path, "wb") as f:
                    f.write(image_data)

                image_paths.append(image_path)
                image_counter += 1

            except Exception as e:
                print(f"Warning: Failed to extract image {image_counter}: {e}")

    return image_paths


def extract_word_content(docx_path: Path) -> Dict[str, Any]:
    """Extract complete structured content from Word document.

    Main extraction function that processes a Word document and extracts:
    - Full text content
    - Paragraph structure with styles
    - Heading hierarchy
    - Images (if any)
    - Raw metadata

    Args:
        docx_path: Path to the .docx file.

    Returns:
        Dictionary containing:
        - raw_text (str): Complete document text
        - paragraphs (List[Dict]): List of paragraph dicts with:
            - index (int): Paragraph index
            - style (str): Word style name
            - text (str): Paragraph text content
            - level (Optional[int]): Heading level (1-9) if heading
            - is_heading (bool): True if paragraph is a heading
        - headings (List[Dict]): List of heading paragraphs only
        - metadata_raw (Dict): Raw metadata from core properties
        - total_paragraphs (int): Total paragraph count
        - has_images (bool): Whether document contains images

    Raises:
        FileNotFoundError: If docx_path does not exist.
        ValueError: If file is not a valid .docx document.

    Example:
        >>> content = extract_word_content(Path("darwin.docx"))
        >>> print(f"Document has {content['total_paragraphs']} paragraphs")
        >>> print(f"Found {len(content['headings'])} headings")
        >>> for h in content['headings']:
        ...     print(f"H{h['level']}: {h['text'][:50]}")
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"Word document not found: {docx_path}")

    if not docx_path.suffix.lower() == ".docx":
        raise ValueError(f"File must be .docx format: {docx_path}")

    # Load document
    doc = Document(docx_path)

    # Extract metadata
    metadata_raw = extract_word_metadata(docx_path)

    # Process paragraphs
    paragraphs: List[Dict[str, Any]] = []
    headings: List[Dict[str, Any]] = []
    full_text_parts: List[str] = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name

        # Determine if this is a heading and its level
        heading_level = _get_heading_level(style_name)
        is_heading = heading_level is not None

        para_dict = {
            "index": idx,
            "style": style_name,
            "text": text,
            "level": heading_level,
            "is_heading": is_heading,
        }

        paragraphs.append(para_dict)

        if is_heading and text:
            headings.append(para_dict)

        # Add to full text (skip empty paragraphs)
        if text:
            full_text_parts.append(text)

    raw_text = "\n\n".join(full_text_parts)

    # Check for images (we'll extract them later if needed)
    has_images = len(doc.part.rels) > 1  # More than just the document.xml relationship

    return {
        "raw_text": raw_text,
        "paragraphs": paragraphs,
        "headings": headings,
        "metadata_raw": metadata_raw,
        "total_paragraphs": len(paragraphs),
        "has_images": has_images,
    }


def build_markdown_from_word(
    paragraphs: List[Dict[str, Any]],
    skip_metadata_lines: int = 5,
) -> str:
    """Build Markdown text from Word document paragraphs.

    Converts Word document structure to Markdown format compatible with
    the existing RAG pipeline. Heading styles are converted to Markdown
    headers (#, ##, ###, etc.).

    Args:
        paragraphs: List of paragraph dicts from extract_word_content().
        skip_metadata_lines: Number of initial paragraphs to skip (metadata).
            Default: 5 (skip TITRE, AUTEUR, EDITION lines).

    Returns:
        Markdown-formatted text.

    Example:
        >>> content = extract_word_content(Path("doc.docx"))
        >>> markdown = build_markdown_from_word(content["paragraphs"])
        >>> with open("output.md", "w") as f:
        ...     f.write(markdown)
    """
    markdown_lines: List[str] = []

    for para in paragraphs[skip_metadata_lines:]:
        text = para["text"]

        if not text:
            continue

        if para["is_heading"] and para["level"]:
            # Convert heading to Markdown: Heading 1 -> #, Heading 2 -> ##, etc.
            level = para["level"]
            markdown_lines.append(f"{'#' * level} {text}")
            markdown_lines.append("")  # Blank line after heading
        else:
            # Normal paragraph
            markdown_lines.append(text)
            markdown_lines.append("")  # Blank line after paragraph

    return "\n".join(markdown_lines).strip()
