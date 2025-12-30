"""Generate Word documents from chat exchanges.

This module provides functionality to export chat conversations between users
and the RAG assistant into formatted Microsoft Word documents (.docx).

Example:
    Export a simple chat exchange::

        from pathlib import Path
        from utils.word_exporter import create_chat_export

        filepath = create_chat_export(
            user_question="What is phenomenology?",
            assistant_response="Phenomenology is a philosophical movement...",
            output_dir=Path("output")
        )

    Export with reformulated question::

        filepath = create_chat_export(
            user_question="What does Husserl mean by phenomenology?",
            assistant_response="Husserl defines phenomenology as...",
            is_reformulated=True,
            original_question="What is phenomenology?",
            output_dir=Path("output")
        )
"""

from pathlib import Path
from typing import Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    raise ImportError(
        "python-docx is required for Word export. "
        "Install with: pip install python-docx"
    )


def create_chat_export(
    user_question: str,
    assistant_response: str,
    is_reformulated: bool = False,
    original_question: Optional[str] = None,
    output_dir: Path = Path("output"),
) -> Path:
    """Create a Word document from a chat exchange.

    Args:
        user_question: The user's question (or reformulated version).
        assistant_response: The assistant's complete response text.
        is_reformulated: Whether the question was reformulated by the system.
            If True, both original and reformulated questions are included.
        original_question: The original user question before reformulation.
            Only used when is_reformulated is True.
        output_dir: Directory where the .docx file will be saved.
            Created if it doesn't exist.

    Returns:
        Path to the generated .docx file.

    Raises:
        OSError: If the output directory cannot be created or file cannot be saved.

    Example:
        >>> from pathlib import Path
        >>> filepath = create_chat_export(
        ...     user_question="Qu'est-ce que la phénoménologie ?",
        ...     assistant_response="La phénoménologie est...",
        ...     output_dir=Path("output")
        ... )
        >>> print(filepath)
        output/chat_export_20250101_123045.docx
    """
    # Create document
    doc = Document()

    # Set margins to 1 inch on all sides
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading("Conversation RAG - Export", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_p.add_run(
        f"Exporté le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    )
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Spacer

    # Original question (if reformulated)
    if is_reformulated and original_question:
        doc.add_heading("Question Originale", level=2)
        original_p = doc.add_paragraph(original_question)
        original_p.style = "Intense Quote"
        doc.add_paragraph()

    # User question
    question_label = "Question Reformulée" if is_reformulated else "Question"
    doc.add_heading(question_label, level=2)
    question_p = doc.add_paragraph(user_question)
    question_p.style = "Intense Quote"

    doc.add_paragraph()  # Spacer

    # Assistant response
    doc.add_heading("Réponse de l'Assistant", level=2)
    response_p = doc.add_paragraph(assistant_response)

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_p.add_run(
        "Généré par Library RAG - Recherche Philosophique"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.italic = True

    # Generate filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"chat_export_{timestamp}.docx"
    filepath = output_dir / filename

    # Save document (create directory if needed)
    output_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(filepath))

    return filepath
